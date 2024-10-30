from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.paginator import Paginator
from .models import Address
from django.contrib import messages
from django.template import Context
from django.template.loader import get_template
from django.http import HttpResponse
from xhtml2pdf import pisa
from docx import Document
from django.db.models import Q
import pandas as pd
from django.contrib.auth.models import User 
import time
from django.contrib.auth import login
from django.db import OperationalError
import logging
from django.core.mail import send_mail
import html 
from .models import MonthlyCopiesSummary, AddressRecord
from .models import Address, AddressRecord, AddressHistory, AddressRecordHistory, MonthlyCopiesSummary  # Import your models
from django.contrib.auth.models import User  # If you need to log the user actions
from django.db.models import Sum

import os
import shutil
from django.conf import settings

from django.contrib.auth.decorators import login_required
from django.core.mail import send_mail
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.contrib import messages
import os
import shutil
from datetime import datetime
from django.conf import settings
from django.shortcuts import redirect
from django.contrib import messages
from .utils import backup_database

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
# Corrected to send_email


from django.shortcuts import render, redirect, get_object_or_404
from .models import Professional, Classified, Expert, Customer, Interview

# Professional Views
def professional_list(request):
    professionals = Professional.objects.all()
    return render(request, 'professional_list.html', {'professionals': professionals})

def professional_create(request):
    if request.method == 'POST':
        customer_name = request.POST.get('customer_name')
        company_name = request.POST.get('company_name')
        designation = request.POST.get('designation')
        edition = request.POST.get('edition')
        Professional.objects.create(
            customer_name=customer_name,
            company_name=company_name,
            designation=designation,
            edition=edition
        )
        return redirect('professional_list')
    return render(request, 'professional_form.html')

def professional_update(request, pk):
    professional = get_object_or_404(Professional, pk=pk)
    if request.method == 'POST':
        professional.customer_name = request.POST.get('customer_name')
        professional.company_name = request.POST.get('company_name')
        professional.designation = request.POST.get('designation')
        professional.edition = request.POST.get('edition')
        professional.save()
        return redirect('professional_list')
    return render(request, 'professional_form.html', {'professional': professional})

def professional_delete(request, pk):
    professional = get_object_or_404(Professional, pk=pk)
    if request.method == 'POST':
        professional.delete()
        return redirect('professional_list')
    return render(request, 'professional_confirm_delete.html', {'professional': professional})

# Classified Views
def classified_list(request):
    classifieds = Classified.objects.all()
    return render(request, 'classified_list.html', {'classifieds': classifieds})

def classified_create(request):
    if request.method == 'POST':
        customer_name = request.POST.get('customer_name')
        company_name = request.POST.get('company_name')
        designation = request.POST.get('designation')
        edition = request.POST.get('edition')
        Classified.objects.create(
            customer_name=customer_name,
            company_name=company_name,
            designation=designation,
            edition=edition
        )
        return redirect('classified_list')
    return render(request, 'classified_form.html')

def classified_update(request, pk):
    classified = get_object_or_404(Classified, pk=pk)
    if request.method == 'POST':
        classified.customer_name = request.POST.get('customer_name')
        classified.company_name = request.POST.get('company_name')
        classified.designation = request.POST.get('designation')
        classified.edition = request.POST.get('edition')
        classified.save()
        return redirect('classified_list')
    return render(request, 'classified_form.html', {'classified': classified})

def classified_delete(request, pk):
    classified = get_object_or_404(Classified, pk=pk)
    if request.method == 'POST':
        classified.delete()
        return redirect('classified_list')
    return render(request, 'classified_confirm_delete.html', {'classified': classified})

# Expert Views
def expert_list(request):
    experts = Expert.objects.all()
    return render(request, 'expert_list.html', {'experts': experts})

def expert_create(request):
    if request.method == 'POST':
        customer_name = request.POST.get('customer_name')
        company_name = request.POST.get('company_name')
        designation = request.POST.get('designation')
        edition = request.POST.get('edition')
        Expert.objects.create(
            customer_name=customer_name,
            company_name=company_name,
            designation=designation,
            edition=edition
        )
        return redirect('expert_list')
    return render(request, 'expert_form.html')

def expert_update(request, pk):
    expert = get_object_or_404(Expert, pk=pk)
    if request.method == 'POST':
        expert.customer_name = request.POST.get('customer_name')
        expert.company_name = request.POST.get('company_name')
        expert.designation = request.POST.get('designation')
        expert.edition = request.POST.get('edition')
        expert.save()
        return redirect('expert_list')
    return render(request, 'expert_form.html', {'expert': expert})

def expert_delete(request, pk):
    expert = get_object_or_404(Expert, pk=pk)
    if request.method == 'POST':
        expert.delete()
        return redirect('expert_list')
    return render(request, 'expert_confirm_delete.html', {'expert': expert})

# Customer Views
def customer_list(request):
    customers = Customer.objects.all()
    return render(request, 'customer_list.html', {'customers': customers})

def customer_create(request):
    if request.method == 'POST':
        customer_name = request.POST.get('customer_name')
        company_name = request.POST.get('company_name')
        designation = request.POST.get('designation')
        edition = request.POST.get('edition')
        Customer.objects.create(
            customer_name=customer_name,
            company_name=company_name,
            designation=designation,
            edition=edition
        )
        return redirect('customer_list')
    return render(request, 'customer_form.html')

def customer_update(request, pk):
    customer = get_object_or_404(Customer, pk=pk)
    if request.method == 'POST':
        customer.customer_name = request.POST.get('customer_name')
        customer.company_name = request.POST.get('company_name')
        customer.designation = request.POST.get('designation')
        customer.edition = request.POST.get('edition')
        customer.save()
        return redirect('customer_list')
    return render(request, 'customer_form.html', {'customer': customer})

def customer_delete(request, pk):
    customer = get_object_or_404(Customer, pk=pk)
    if request.method == 'POST':
        customer.delete()
        return redirect('customer_list')
    return render(request, 'customer_confirm_delete.html', {'customer': customer})

# Interview Views
def interview_list(request):
    interviews = Interview.objects.all()
    return render(request, 'interview_list.html', {'interviews': interviews})

def interview_create(request):
    if request.method == 'POST':
        customer_name = request.POST.get('customer_name')
        company_name = request.POST.get('company_name')
        designation = request.POST.get('designation')
        edition = request.POST.get('edition')
        Interview.objects.create(
            customer_name=customer_name,
            company_name=company_name,
            designation=designation,
            edition=edition
        )
        return redirect('interview_list')
    return render(request, 'interview_form.html')

def interview_update(request, pk):
    interview = get_object_or_404(Interview, pk=pk)
    if request.method == 'POST':
        interview.customer_name = request.POST.get('customer_name')
        interview.company_name = request.POST.get('company_name')
        interview.designation = request.POST.get('designation')
        interview.edition = request.POST.get('edition')
        interview.save()
        return redirect('interview_list')
    return render(request, 'interview_form.html', {'interview': interview})

def interview_delete(request, pk):
    interview = get_object_or_404(Interview, pk=pk)
    if request.method == 'POST':
        interview.delete()
        return redirect('interview_list')
    return render(request, 'interview_confirm_delete.html', {'interview': interview})

def magazine_details(request):
    return render(request, 'magazine_details.html')




@login_required
def backup_option(request):
    if request.method == 'POST':
        # Check the value of the 'backup' button
        if request.POST.get('backup') == 'yes':
            db_path = 'C:/Users/BrandsMart/Downloads/dj (2)/dj/dj/address/myproject/db.sqlite3'
            backup_dir = 'C:/Users/BrandsMart/Documents/backups'
            
            # Create backup
            backup_file = backup_database(db_path, backup_dir)
            if backup_file:
                # Redirect to address_list after successful backup
                return redirect('address_list')
            else:
                return render(request, 'error.html', {'message': 'Backup failed.'})

        elif request.POST.get('backup') == 'no':
            # Redirect to address_list if the user chooses not to back up
            return redirect('address_list')

    return render(request, 'backup_option.html')




def register(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        password_confirm = request.POST['password_confirm']
        is_admin = request.POST.get('is_admin') == 'True'

        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already exists.")
        elif password != password_confirm:
            messages.error(request, "Passwords do not match.")
        else:
            try:
                user = User.objects.create_user(username=username, password=password)
                user.is_staff = is_admin
                user.save()
                messages.success(request, 'Registration successful! You can now log in.')
                login(request, user)  # Automatically log in the user after registration
                return redirect('login')  # Redirect to the address list
            except Exception as e:
                messages.error(request, str(e))

    return render(request, 'registration/register.html')



@login_required
def address_list(request):
    search_field = request.GET.get('field', '')  # Get the selected field from the form
    search_query = request.GET.get('search', '')  # Get the search query from the form

    # Base queryset
    addresses = Address.objects.all()

    # Apply search filter
    if search_query:
        if search_field == 's_no':
            addresses = addresses.filter(s_no__icontains=search_query)
        elif search_field == 'region':
            addresses = addresses.filter(region__icontains=search_query)
        elif search_field == 'categories':
            addresses = addresses.filter(categories__icontains=search_query)
        elif search_field == 'company_name':
            addresses = addresses.filter(company_name__icontains=search_query)
        elif search_field == 'postal_dtdc':
            addresses = addresses.filter(postal_dtdc__icontains=search_query)
        elif search_field == 'person_name':
            addresses = addresses.filter(person_name__icontains=search_query)
        elif search_field == 'address':
            addresses = addresses.filter(address__icontains=search_query)
        elif search_field == 'phone_number':
            addresses = addresses.filter(phone_number__icontains=search_query)
        elif search_field == 'email_id':
            addresses = addresses.filter(email_id__icontains=search_query)

    # Check if the user is an admin (staff)
    is_admin = request.user.is_staff

    # Calculate the total copies sum
    total_copies = addresses.aggregate(total_copies=Sum('copies'))['total_copies'] or 0

    return render(request, 'addresses/address_list.html', {
        'addresses': addresses,
        'is_admin': is_admin,
        'search_query': search_query,
        'search_field': search_field,
        'total_copies': total_copies,
       
    })

def download_pdf(request):
    selected_ids = request.POST.getlist('selected_addresses')

    if not selected_ids:
        return HttpResponse('No addresses selected to download.')

    addresses = Address.objects.filter(pk__in=selected_ids)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="address_list.pdf"'

    # Create PDF logic
    template_path = 'addresses/print_address_list.html'  # Adjust as necessary
    context = {'addresses': addresses}
    template = render(request, template_path, context)
    pisa_status = pisa.CreatePDF(template, dest=response)

    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html.escape(pisa_status.err) + '</pre>')

    return response

from io import BytesIO
from django.http import HttpResponse
from django.shortcuts import render
from docx import Document
from .models import Address

def download_word(request):
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_addresses')

        if not selected_ids:
            return HttpResponse('No addresses selected to download.')

        # Fetch addresses based on selected IDs
        addresses = Address.objects.all() if 'all' in selected_ids else Address.objects.filter(pk__in=selected_ids)

        # Create a Word document
        doc = Document()
        doc.add_heading('Address List', level=1)

        # Add a table to the document
        table = doc.add_table(rows=1, cols=4)  # 4 columns as required
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'S.No'
        hdr_cells[1].text = 'Region'
        hdr_cells[2].text = 'Postal/DTDC'
        hdr_cells[3].text = 'Details'  # This will include person name, company name, address, phone number

        for address in addresses:
            row_cells = table.add_row().cells
            row_cells[0].text = str(address.s_no)
            row_cells[1].text = address.region
            row_cells[2].text = address.postal_dtdc

            # Combine person name, company name, address, and phone number into one column
            details = (
                f"Name: {address.person_name}\n"
                f"Company: {address.company_name}\n"
                f"Address: {address.address}\n"
                f"Phone: {address.phone_number}"
            )
            row_cells[3].text = details

        # Create a BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Create response
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="address_list.docx"'

        return response

    # Render the form if not a POST request
    addresses = Address.objects.all()  # Fetch all addresses for the selection
    return render(request, 'download_form.html', {'addresses': addresses})


   


def welcome(request):
    return render(request, 'welcome.html')

def is_admin(user):
    return user.is_admin


def print_address_list(request):
    addresses = Address.objects.all()  # Fetch all addresses
    return render(request, 'addresses/print_address_list.html', {'addresses': addresses})
def custom_login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            # Check if the user is an admin
            if user.is_staff:
                return redirect('backup_option')  # Redirect to the backup option page
            else:
                return redirect('home')  # Redirect to a normal user home page
        else:
            return render(request, 'registration/login.html', {'error': 'Invalid credentials'})

    return render(request, 'registration/login.html')


@login_required
def address_create(request):
    if request.method == 'POST':
        # Collecting data from the POST request
        region = request.POST.get('region')
        categories = request.POST.get('categories')
        postal_dtdc = request.POST.get('postal_dtdc')
        person_name = request.POST.get('person_name')
        company_name = request.POST.get('company_name')
        address = request.POST.get('address')
        phone_number = request.POST.get('phone_number')
        copies = request.POST.get('copies')

        # Convert to integer only if it's not None
        if copies:
            try:
                copies = int(copies)
            except ValueError:
                print("Error converting copies to integer")
                copies = 1

        receiver_name = request.POST.get('receiver_name')
        email_id = request.POST.get('email_id')

        # Saving the data to the database
        new_address = Address(
            region=region,
            categories=categories,
            postal_dtdc=postal_dtdc,
            person_name=person_name,
            company_name=company_name,
            address=address,
            phone_number=phone_number,
            copies=copies,
            receiver_name=receiver_name,
            email_id=email_id
        )
        new_address.save()
        
        # Log action with the correct ID
        AddressHistory.objects.create(
            user=request.user,
            action='create',
            address_record=new_address
        )

        # Display a success message
        messages.success(request, 'Address added successfully!')

        # Redirect to address list
        return redirect('address_list')

    return render(request, 'addresses/address_form.html')


def address_update(request, pk):
    address = get_object_or_404(Address, pk=pk)

    if request.method == 'POST':
        # Update the address fields directly from POST data
        address.region = request.POST.get('region', address.region)
        address.categories = request.POST.get('categories', address.categories)
        address.postal_dtdc = request.POST.get('postal_dtdc', address.postal_dtdc)
        address.person_name = request.POST.get('person_name', address.person_name)
        address.company_name = request.POST.get('company_name', address.company_name)
        address.address = request.POST.get('address', address.address)
        address.phone_number = request.POST.get('phone_number', address.phone_number)
        address.copies = request.POST.get('copies', address.copies)
        address.receiver_name = request.POST.get('receiver_name', address.receiver_name)
        address.email_id = request.POST.get('email_id', address.email_id)
        
        address.save()  # Save the updated address
        
        AddressHistory.objects.create(
            user=request.user,
            action='edit',
            address_record=address
        )
        messages.success(request, "Address updated successfully.")
        return redirect('address_list')  # Adjust to your actual address list view

    return render(request, 'addresses/update_form.html', {'address': address})


def address_delete(request, pk):
    address = get_object_or_404(Address, pk=pk)

    if request.method == 'POST':
        address.delete()
        AddressHistory.objects.create(
            user=request.user,
            action='delete',
            address_record=address
        )
        messages.success(request, "Address deleted successfully.")
        return redirect('address_list')  # Adjust to your actual address list view

    return render(request, 'addresses/confirm_delete.html', {'address': address})



@login_required
def import_addresses(request):
    if request.method == 'POST':
        excel_file = request.FILES.get('file')
        if excel_file:
            # Check if the uploaded file is an Excel file
            if not excel_file.name.endswith(('.xls', '.xlsx')):
                messages.error(request, "Please upload a valid Excel file.")
                return redirect('import_form')  # Use the correct URL name for your import form

            try:
                # Read the Excel file
                df = pd.read_excel(excel_file)

                # Initialize counter for imported addresses
                success_count = 0

                # Loop through the rows in the DataFrame
                for index, row in df.iterrows():
                    address_value = row['address']

                    # Delete existing duplicates based on the address field
                    Address.objects.filter(address=address_value).delete()

                    try:
                        # Create a new Address object (after deleting any existing duplicates)
                        Address.objects.create(
                            region=row['region'],
                            categories=row['categories'],
                            postal_dtdc=row['postal_dtdc'],
                            person_name=row['person_name'],
                            company_name=row['company_name'],
                            address=address_value,
                            phone_number=row['phone_number'],
                            copies=row['copies'],
                            receiver_name=row['receiver_name'],
                            email_id=row['email_id'],
                        )
                        success_count += 1
                    except Exception as e:
                        # Log or handle errors as needed
                        continue

                # Provide user feedback about the import process
                messages.success(request, f"Imported {success_count} addresses successfully after handling duplicates.")
                return redirect('address_list')  # Redirect to the address list after import

            except Exception as e:
                messages.error(request, "Error reading the Excel file. Please check the file format.")
                return redirect('import_form')  # Redirect to the import form in case of error

    return render(request, 'addresses/import_form.html')  # Render the import form template



@login_required
def return_address_list(request):
    records = AddressRecord.objects.all()
    return render(request, 'return_address_list.html', {'records': records})

@login_required
def return_create_address(request):
    if request.method == 'POST':
        s_no = request.POST['s_no']
        return_address = request.POST['return_address']
        confirm_address = request.POST['confirm_address']
        reason = request.POST['reason']
        status = request.POST.get('status', 'not_verified')

        # Create the AddressRecord instance
        new_record = AddressRecord.objects.create(
            s_no=s_no,
            return_address=return_address,
            confirm_address=confirm_address,
            reason=reason,
            status=status
        )

        # Log the action
        AddressRecordHistory.objects.create(
            user=request.user,
            action='create',
            address_record=new_record
        )
        return redirect('return_address_list')

    return render(request, 'return_create_address.html')

@login_required
def return_edit_address(request, pk):
    record = get_object_or_404(AddressRecord, pk=pk)

    if request.method == 'POST':
        record.s_no = request.POST['s_no']
        record.return_address = request.POST['return_address']
        record.confirm_address = request.POST['confirm_address']
        record.reason = request.POST['reason']
        record.status = request.POST['status']
        record.save()

        # Log the action
        AddressRecordHistory.objects.create(
            user=request.user,
            action='edit',
            address_record=record
        )
        return redirect('return_address_list')

    return render(request, 'return_edit_address.html', {'record': record})

@login_required
def return_delete_address(request, pk):
    record = get_object_or_404(AddressRecord, pk=pk)
    
    # Log the action before deleting
    AddressRecordHistory.objects.create(
            user=request.user,
            action='delete',
            address_record=record
        )
   
    record.delete()
    return redirect('return_address_list')


def monthly_copies_summary(request):
    summaries = MonthlyCopiesSummary.objects.all()
    month_names = {
        1: 'January', 2: 'February', 3: 'March',
        4: 'April', 5: 'May', 6: 'June',
        7: 'July', 8: 'August', 9: 'September',
        10: 'October', 11: 'November', 12: 'December'
    }

    # Convert start and end month numbers to names
    for summary in summaries:
        summary.start_month_name = month_names.get(summary.start_month, summary.start_month)
        summary.end_month_name = month_names.get(summary.end_month, summary.end_month)

    return render(request, 'monthly_copies_summary.html', {'summaries': summaries})


# # List View for Monthly Copies Summary
# def monthly_copies_summary(request):
#     summaries = MonthlyCopiesSummary.objects.all()
#     return render(request, 'monthly_copies_summary.html', {'summaries': summaries})


# # Create View for Monthly Copies Summary
# def create_monthly_copies_summary(request):
#     if request.method == 'POST':
#         year = request.POST.get('year')
#         start_month = request.POST.get('start_month')
#         end_month = request.POST.get('end_month')
#         save_copies = request.POST.get('save_copies') == 'on'

#         month_mapping = {
#             'january': 1, 'february': 2, 'march': 3,
#             'april': 4, 'may': 5, 'june': 6,
#             'july': 7, 'august': 8, 'september': 9,
#             'october': 10, 'november': 11, 'december': 12
#         }

#         start_month_num = month_mapping.get(start_month.lower())
#         end_month_num = month_mapping.get(end_month.lower())

#         if start_month_num and end_month_num:
#             total_copies = Address.objects.aggregate(Sum('copies'))['copies__sum'] or 0

#             MonthlyCopiesSummary.objects.create(
#                 year=year,
#                 start_month=start_month_num,
#                 end_month=end_month_num,
#                 total_copies=total_copies
#             )
#             return redirect('monthly_copies_summary_list')

#     total_copies = Address.objects.aggregate(Sum('copies'))['copies__sum'] or 0
#     return render(request, 'create_monthly_copies_summary.html', {
#         'total_copies': total_copies
#     })

from django.shortcuts import render, redirect
from django.db.models import Sum
from .models import MonthlyCopiesSummary, Address  # Adjust imports based on your models

# Create View for Monthly Copies Summary
def create_monthly_copies_summary(request):
    if request.method == 'POST':
        year = request.POST.get('year')
        start_month = request.POST.get('start_month')
        end_month = request.POST.get('end_month')
        save_copies = request.POST.get('save_copies') == 'on'

        month_mapping = {
            'january': 1, 'february': 2, 'march': 3,
            'april': 4, 'may': 5, 'june': 6,
            'july': 7, 'august': 8, 'september': 9,
            'october': 10, 'november': 11, 'december': 12,
            'jan': 1, 'feb': 2, 'mar': 3,
            'apr': 4, 'jun': 6, 'jul': 7,
            'aug': 8, 'sep': 9, 'oct': 10,
            'nov': 11, 'dec': 12
        }

        # Convert start and end months to lowercase and find their mapping
        start_month_num = month_mapping.get(start_month.lower())
        end_month_num = month_mapping.get(end_month.lower())

        # Check if months are valid
        if start_month_num is None or end_month_num is None:
            return render(request, 'create_monthly_copies_summary.html', {
                'error': 'Invalid month names provided.',
                'total_copies': Address.objects.aggregate(Sum('copies'))['copies__sum'] or 0
            })

        # Calculate total copies
        total_copies = Address.objects.aggregate(Sum('copies'))['copies__sum'] or 0

        # Create Monthly Copies Summary entry
        MonthlyCopiesSummary.objects.create(
            year=year,
            start_month=start_month_num,
            end_month=end_month_num,
            total_copies=total_copies
        )
        return redirect('monthly_copies_summary')

    total_copies = Address.objects.aggregate(Sum('copies'))['copies__sum'] or 0
    return render(request, 'create_monthly_copies_summary.html', {
        'total_copies': total_copies
    })



# Edit View for Monthly Copies Summary
def edit_monthly_copies_summary(request, id):
    summary = get_object_or_404(MonthlyCopiesSummary, id=id)

    if request.method == 'POST':
        summary.year = request.POST.get('year')
        summary.start_month = request.POST.get('start_month')
        summary.end_month = request.POST.get('end_month')
        summary.save()
        return redirect('monthly_copies_summary')

    return render(request, 'edit_monthly_copies_summary.html', {
        'summary': summary,
    })

# Delete View for Monthly Copies Summary
def delete_monthly_copies_summary(request, id):
    summary = get_object_or_404(MonthlyCopiesSummary, id=id)

    if request.method == 'POST':
        summary.delete()
        return redirect('monthly_copies_summary')

    return render(request, 'delete_monthly_copies_summary.html', {
        'summary': summary,
    })


def save_monthly_copies(request):
    # Your code here
    pass



# views.py

@login_required
def return_address_history(request):
    history_list = AddressRecordHistory.objects.all()  # Fetch all history records
    return render(request, 'return_history_list.html', {'history_list': history_list})

@login_required
def create_address_history(request):
    history_list = AddressHistory.objects.all()
    return render(request, 'create_history.html', {'history_list': history_list})
